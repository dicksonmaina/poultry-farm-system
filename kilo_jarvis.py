"""
╔══════════════════════════════════════════════════════════════╗
║          KILO JARVIS v2.0 — Telegram AI Assistant            ║
║   Features: Image Analysis, PDF/DOC, Web Search, Memory      ║
║   Author: Richard Dickson (dicksonmaina) — Nairobi, Kenya     ║
╚══════════════════════════════════════════════════════════════╝

Install dependencies:
    pip install requests reportlab python-docx
"""

import requests
import time
import json
import logging
import os
import base64
import re
from datetime import datetime

# ── Optional imports ───────────────────────────────────────────
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                    HRFlowable)
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("⚠️  reportlab missing → pip install reportlab")

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️  python-docx missing → pip install python-docx")

# ── Logging ────────────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler("jarvis.log", encoding="utf-8")
    ]
)
log = logging.getLogger(__name__)

# ══════════════════════════════════════════════════════════════
#  CONFIG  ← Replace TELEGRAM_TOKEN with your bot token
# ══════════════════════════════════════════════════════════════
TELEGRAM_TOKEN    = "YOUR_TELEGRAM_BOT_TOKEN"
GROQ_API_KEY      = "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"
GROQ_MODEL        = "llama-3.1-8b-instant"
GROQ_MODEL_VISION = "llama-3.2-11b-vision-preview"
MAX_TOKENS        = 1024
POLL_TIMEOUT      = 30
BOT_OWNER         = "Richard"

TELEGRAM_BASE = f"https://api.telegram.org/bot{TELEGRAM_TOKEN}"
GROQ_URL      = "https://api.groq.com/openai/v1/chat/completions"

SYSTEM_PROMPT = f"""You are JARVIS, an advanced AI assistant for {BOT_OWNER} (dicksonmaina),
a CS diploma student at NIBS Technical College, Nairobi Kenya.

Personality:
- Smart, direct, witty like Iron Man's JARVIS
- You know Richard's projects: Poultry Farm Management System, JARVIS OS ecosystem
- His stack: PHP, Python, Supabase, Tailwind CSS, Alpine.js, Ollama, n8n, Neo4j
- Give practical, code-ready answers
- Kenyan context awareness: M-Pesa, Safaricom Daraja, local market

Capabilities you have:
- Generate PDF reports (/pdf [topic])
- Generate Word docs (/doc [topic])
- Analyze images sent by the user
- Remember conversation context (last 20 messages)

Always be helpful, accurate, and encouraging."""

# ══════════════════════════════════════════════════════════════
#  CONVERSATION MEMORY
# ══════════════════════════════════════════════════════════════
conversation_history = {}
user_stats = {}

def get_history(chat_id):
    if chat_id not in conversation_history:
        conversation_history[chat_id] = []
    return conversation_history[chat_id]

def add_to_history(chat_id, role, content):
    history = get_history(chat_id)
    history.append({"role": role, "content": content})
    if len(history) > 20:
        conversation_history[chat_id] = history[-20:]

def track_usage(chat_id, user_name):
    if chat_id not in user_stats:
        user_stats[chat_id] = {
            "name": user_name,
            "messages": 0,
            "first_seen": datetime.now().strftime("%Y-%m-%d")
        }
    user_stats[chat_id]["messages"] += 1

# ══════════════════════════════════════════════════════════════
#  GROQ — TEXT
# ══════════════════════════════════════════════════════════════
def ask_ai(message, history=[]):
    import requests, os
    nim_key = os.environ.get("NVIDIA_API_KEY","")
    groq_key = os.environ.get("GROQ_API_KEY","")
    
    if nim_key:
        try:
            r = requests.post(
                "https://integrate.api.nvidia.com/v1/chat/completions",
                headers={"Authorization":f"Bearer {nim_key}","Content-Type":"application/json"},
                json={"model":"moonshotai/kimi-k2-thinking",
                      "messages":[{"role":"user","content":message}],
                      "max_tokens":500},
                timeout=15
            )
            if r.status_code==200:
                return r.json()["choices"][0]["message"]["content"]
        except: pass
    
    if groq_key:
        r = requests.post(
            "https://api.groq.com/openai/v1/chat/completions",
            headers={"Authorization":f"Bearer {groq_key}","Content-Type":"application/json"},
            json={"model":"llama-3.1-8b-instant",
                  "messages":[{"role":"user","content":message}],
                  "max_tokens":500},
            timeout=15
        )
        return r.json()["choices"][0]["message"]["content"]
    
    return "No provider available"

# ══════════════════════════════════════════════════════════════
#  GROQ — VISION (Image Analysis)
# ══════════════════════════════════════════════════════════════
def call_groq_vision(chat_id, image_b64, user_prompt="Analyze this image in detail."):
    messages = [
        {"role": "system", "content": SYSTEM_PROMPT},
        {
            "role": "user",
            "content": [
                {
                    "type": "image_url",
                    "image_url": {"url": f"data:image/jpeg;base64,{image_b64}"}
                },
                {"type": "text", "text": user_prompt}
            ]
        }
    ]

    headers = {
        "Authorization": f"Bearer {GROQ_API_KEY}",
        "Content-Type": "application/json"
    }
    payload = {
        "model": GROQ_MODEL_VISION,
        "messages": messages,
        "max_tokens": 1024
    }

    for attempt in range(3):
        try:
            r = requests.post(GROQ_URL, headers=headers, json=payload, timeout=30)
            if r.status_code == 200:
                reply = r.json()["choices"][0]["message"]["content"].strip()
                add_to_history(chat_id, "assistant", f"[Analyzed image]: {reply}")
                return reply
            elif r.status_code == 429:
                wait = (2 ** attempt) * 4
                log.warning(f"Vision rate limited. Waiting {wait}s")
                time.sleep(wait)
            elif r.status_code in (400, 404):
                # Vision model unavailable — fallback
                msg = (f"The user sent an image with caption: '{user_prompt}'. "
                       "The vision model is currently unavailable. "
                       "Acknowledge this politely and ask them to describe the image in text.")
                add_to_history(chat_id, "user", msg)
                response = ask_ai(msg)
                add_to_history(chat_id, "assistant", response)
                return response
            else:
                log.error(f"Vision {r.status_code}: {r.text[:200]}")
                return f"⚠️ Image analysis failed ({r.status_code}). Describe it in text instead."
        except Exception as e:
            log.error(f"Vision error: {e}")
            time.sleep(3)

    return "⚠️ Could not analyze image. Please describe what you need help with."

# ══════════════════════════════════════════════════════════════
#  TELEGRAM HELPERS
# ══════════════════════════════════════════════════════════════
def send_message(chat_id, text, parse_mode="Markdown"):
    if len(text) > 4000:
        chunks = [text[i:i+4000] for i in range(0, len(text), 4000)]
        for chunk in chunks:
            _send_single(chat_id, chunk, parse_mode)
            time.sleep(0.3)
    else:
        _send_single(chat_id, text, parse_mode)

def _send_single(chat_id, text, parse_mode="Markdown"):
    try:
        r = requests.post(
            f"{TELEGRAM_BASE}/sendMessage",
            json={"chat_id": chat_id, "text": text, "parse_mode": parse_mode},
            timeout=10
        )
        if r.status_code != 200:
            # Retry without markdown if formatting broke
            if "parse" in r.text.lower() or "entities" in r.text.lower():
                requests.post(
                    f"{TELEGRAM_BASE}/sendMessage",
                    json={"chat_id": chat_id, "text": text},
                    timeout=10
                )
            else:
                log.error(f"Send failed: {r.text[:100]}")
    except Exception as e:
        log.error(f"Send error: {e}")

def send_document(chat_id, filepath, caption=""):
    try:
        with open(filepath, "rb") as f:
            r = requests.post(
                f"{TELEGRAM_BASE}/sendDocument",
                data={"chat_id": chat_id, "caption": caption, "parse_mode": "Markdown"},
                files={"document": f},
                timeout=30
            )
        if r.status_code == 200:
            log.info(f"Document sent: {filepath}")
            return True
        log.error(f"Doc send failed: {r.text[:150]}")
        return False
    except Exception as e:
        log.error(f"Send doc error: {e}")
        return False

def send_typing(chat_id):
    try:
        requests.post(
            f"{TELEGRAM_BASE}/sendChatAction",
            json={"chat_id": chat_id, "action": "typing"},
            timeout=5
        )
    except:
        pass

def send_upload_action(chat_id):
    try:
        requests.post(
            f"{TELEGRAM_BASE}/sendChatAction",
            json={"chat_id": chat_id, "action": "upload_document"},
            timeout=5
        )
    except:
        pass

def download_telegram_file(file_id):
    """Download Telegram file → base64 string"""
    try:
        r = requests.get(f"{TELEGRAM_BASE}/getFile?file_id={file_id}", timeout=10)
        if r.status_code != 200:
            return None
        file_path = r.json()["result"]["file_path"]
        url = f"https://api.telegram.org/file/bot{TELEGRAM_TOKEN}/{file_path}"
        r2 = requests.get(url, timeout=20)
        if r2.status_code == 200:
            return base64.b64encode(r2.content).decode("utf-8")
        return None
    except Exception as e:
        log.error(f"Download error: {e}")
        return None

def get_updates(offset=None):
    params = {
        "timeout": POLL_TIMEOUT,
        "allowed_updates": ["message"]
    }
    if offset:
        params["offset"] = offset
    try:
        r = requests.get(
            f"{TELEGRAM_BASE}/getUpdates",
            params=params,
            timeout=POLL_TIMEOUT + 5
        )
        if r.status_code == 200:
            return r.json().get("result", [])
        log.error(f"getUpdates {r.status_code}")
        time.sleep(3)
        return []
    except requests.exceptions.Timeout:
        return []
    except Exception as e:
        log.error(f"getUpdates error: {e}")
        time.sleep(5)
        return []

# ══════════════════════════════════════════════════════════════
#  PDF GENERATOR
# ══════════════════════════════════════════════════════════════
def generate_pdf(topic, content, filename=None):
    if not PDF_AVAILABLE:
        return None, "reportlab not installed. Run: pip install reportlab"

    if not filename:
        safe = re.sub(r'[^a-zA-Z0-9_]', '_', topic[:30])
        filename = f"/tmp/jarvis_{safe}_{int(time.time())}.pdf"

    try:
        doc = SimpleDocTemplate(
            filename, pagesize=A4,
            rightMargin=2*cm, leftMargin=2*cm,
            topMargin=2*cm, bottomMargin=2*cm
        )
        styles = getSampleStyleSheet()

        # Custom styles
        s_title = ParagraphStyle('JTitle', parent=styles['Title'],
            fontSize=22, textColor=colors.HexColor('#1a1a2e'),
            spaceAfter=6, fontName='Helvetica-Bold', alignment=TA_CENTER)

        s_sub = ParagraphStyle('JSub', parent=styles['Normal'],
            fontSize=10, textColor=colors.HexColor('#6c757d'),
            spaceAfter=20, alignment=TA_CENTER)

        s_h2 = ParagraphStyle('JH2', parent=styles['Heading2'],
            fontSize=13, textColor=colors.HexColor('#0d6efd'),
            spaceBefore=14, spaceAfter=6, fontName='Helvetica-Bold')

        s_h3 = ParagraphStyle('JH3', parent=s_h2,
            fontSize=11, textColor=colors.HexColor('#495057'))

        s_body = ParagraphStyle('JBody', parent=styles['Normal'],
            fontSize=11, leading=16, spaceAfter=8,
            textColor=colors.HexColor('#212529'), alignment=TA_JUSTIFY)

        s_code = ParagraphStyle('JCode', parent=styles['Code'],
            fontSize=9, backColor=colors.HexColor('#f8f9fa'),
            borderColor=colors.HexColor('#dee2e6'), borderWidth=1,
            borderPadding=8, fontName='Courier', spaceAfter=10)

        s_bullet = ParagraphStyle('JBullet', parent=s_body,
            leftIndent=20, firstLineIndent=-10)

        s_footer = ParagraphStyle('JFooter', parent=styles['Normal'],
            fontSize=8, textColor=colors.HexColor('#adb5bd'), alignment=TA_CENTER)

        story = []

        # Header
        story.append(Paragraph(topic, s_title))
        story.append(Paragraph(
            f"Generated by JARVIS AI  •  {datetime.now().strftime('%B %d, %Y at %H:%M')}",
            s_sub
        ))
        story.append(HRFlowable(width="100%", thickness=2,
                                color=colors.HexColor('#0d6efd'), spaceAfter=16))

        # Content parser
        lines = content.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].strip()

            if not line:
                story.append(Spacer(1, 6))
            elif line.startswith('# '):
                story.append(Paragraph(line[2:], s_title))
            elif line.startswith('## '):
                story.append(Paragraph(line[3:], s_h2))
            elif line.startswith('### '):
                story.append(Paragraph(line[4:], s_h3))
            elif line.startswith('```'):
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                story.append(Paragraph(
                    '\n'.join(code_lines).replace('\n', '<br/>'), s_code
                ))
            elif line.startswith(('- ', '* ')):
                story.append(Paragraph(f"• {line[2:]}", s_bullet))
            elif re.match(r'^\d+\.', line):
                story.append(Paragraph(line, s_body))
            else:
                clean = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', line)
                clean = re.sub(r'\*(.*?)\*', r'<i>\1</i>', clean)
                clean = re.sub(r'`(.*?)`', r'<font name="Courier">\1</font>', clean)
                story.append(Paragraph(clean, s_body))
            i += 1

        # Footer
        story.append(Spacer(1, 20))
        story.append(HRFlowable(width="100%", thickness=1,
                                color=colors.HexColor('#dee2e6'), spaceAfter=8))
        story.append(Paragraph(
            "JARVIS AI  •  Built by Richard Dickson  •  Nairobi, Kenya",
            s_footer
        ))

        doc.build(story)
        return filename, None

    except Exception as e:
        log.error(f"PDF error: {e}")
        return None, str(e)

# ══════════════════════════════════════════════════════════════
#  WORD DOC GENERATOR
# ══════════════════════════════════════════════════════════════
def generate_docx(topic, content, filename=None):
    if not DOCX_AVAILABLE:
        return None, "python-docx not installed. Run: pip install python-docx"

    if not filename:
        safe = re.sub(r'[^a-zA-Z0-9_]', '_', topic[:30])
        filename = f"/tmp/jarvis_{safe}_{int(time.time())}.docx"

    try:
        doc = Document()

        # Margins
        for section in doc.sections:
            section.top_margin    = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin   = Inches(1.2)
            section.right_margin  = Inches(1.2)

        # Title
        t = doc.add_heading(topic, 0)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in t.runs:
            run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
            run.font.size = Pt(22)

        # Subtitle
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = sub.add_run(
            f"Generated by JARVIS AI  •  {datetime.now().strftime('%B %d, %Y at %H:%M')}"
        )
        r.font.color.rgb = RGBColor(0x6c, 0x75, 0x7d)
        r.font.size = Pt(10)
        doc.add_paragraph()

        # Content parser
        lines = content.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].strip()

            if not line:
                doc.add_paragraph()
            elif line.startswith('# '):
                h = doc.add_heading(line[2:], 1)
                for r in h.runs:
                    r.font.color.rgb = RGBColor(0x0d, 0x6e, 0xfd)
            elif line.startswith('## '):
                h = doc.add_heading(line[3:], 2)
                for r in h.runs:
                    r.font.color.rgb = RGBColor(0x0d, 0x6e, 0xfd)
            elif line.startswith('### '):
                doc.add_heading(line[4:], 3)
            elif line.startswith('```'):
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                p = doc.add_paragraph('\n'.join(code_lines))
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
            elif line.startswith(('- ', '* ')):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif re.match(r'^\d+\.', line):
                doc.add_paragraph(line, style='List Number')
            else:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', line)
                for part in parts:
                    if not part:
                        continue
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                        run.font.size = Pt(11)
                    elif part.startswith('*') and part.endswith('*'):
                        run = p.add_run(part[1:-1])
                        run.italic = True
                        run.font.size = Pt(11)
                    elif part.startswith('`') and part.endswith('`'):
                        run = p.add_run(part[1:-1])
                        run.font.name = 'Courier New'
                        run.font.size = Pt(10)
                    else:
                        run = p.add_run(part)
                        run.font.size = Pt(11)
            i += 1

        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph(
            "JARVIS AI  •  Built by Richard Dickson  •  Nairobi, Kenya"
        )
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0xad, 0xb5, 0xbd)

        doc.save(filename)
        return filename, None

    except Exception as e:
        log.error(f"DOCX error: {e}")
        return None, str(e)

# ══════════════════════════════════════════════════════════════
#  DOCUMENT GENERATION FLOW
# ══════════════════════════════════════════════════════════════
def handle_generate_doc(chat_id, topic, doc_type="pdf"):
    send_typing(chat_id)
    send_message(chat_id, f"⚙️ Generating *{doc_type.upper()}*: _{topic}_\nPlease wait...")

    prompt = (
        f"Write a comprehensive, well-structured document about: {topic}\n\n"
        "Format using markdown:\n"
        "- Use ## for main sections\n"
        "- Use ### for subsections\n"
        "- Use bullet points and numbered lists\n"
        "- Include code examples if technical\n"
        "- Be thorough and practical\n"
        "- Minimum 400 words"
    )
    system = (
        "You are a professional technical writer. Generate well-structured, "
        "comprehensive documents in markdown format. Include practical examples."
    )

    add_to_history(chat_id, "user", prompt)
    full_prompt = f"{system}\n\n{prompt}"
    content = ask_ai(full_prompt)
    if content and not content.startswith(("⚠️", "❌")):
        add_to_history(chat_id, "assistant", content)

    if not content or content.startswith(("⚠️", "❌")):
        send_message(chat_id, f"❌ Could not generate content:\n{content}")
        return

    send_upload_action(chat_id)

    if doc_type == "pdf":
        if not PDF_AVAILABLE:
            send_message(chat_id,
                "⚠️ *reportlab not installed*\n"
                "Run: `pip install reportlab`\n\n"
                "Here's the content as text:\n\n" + content[:3000])
            return
        filepath, err = generate_pdf(topic, content)
    else:
        if not DOCX_AVAILABLE:
            send_message(chat_id,
                "⚠️ *python-docx not installed*\n"
                "Run: `pip install python-docx`\n\n"
                "Here's the content as text:\n\n" + content[:3000])
            return
        filepath, err = generate_docx(topic, content)

    if filepath and not err:
        sent = send_document(chat_id, filepath,
                             caption=f"📄 *{topic}*\n_Generated by JARVIS AI_")
        if sent:
            send_message(chat_id, f"✅ Your {doc_type.upper()} is ready!")
        else:
            send_message(chat_id, "⚠️ File created but failed to send. Check permissions.")
        try:
            os.remove(filepath)
        except:
            pass
    else:
        send_message(chat_id,
            f"❌ {doc_type.upper()} generation failed: {err}\n\n"
            f"Content preview:\n\n{content[:2000]}")

# ══════════════════════════════════════════════════════════════
#  COMMAND HANDLERS
# ══════════════════════════════════════════════════════════════
def handle_start(chat_id, user_name):
    send_message(chat_id,
        f"⚡ *JARVIS v2.0 — Online*\n"
        f"Welcome, {user_name}!\n\n"
        "━━━━━━━━━━━━━━━━━━━\n"
        "🤖 *What I can do:*\n\n"
        "💬 *Chat* — Ask me anything, I remember context\n"
        "🖼 *Images* — Send photo + optional caption\n"
        "📄 `/pdf [topic]` — Generate PDF document\n"
        "📝 `/doc [topic]` — Generate Word document\n"
        "💻 *Code* — Debug, review, write code\n\n"
        "━━━━━━━━━━━━━━━━━━━\n"
        "⌨️ *All Commands:*\n"
        "`/start` — This menu\n"
        "`/pdf [topic]` — Generate PDF\n"
        "`/doc [topic]` — Generate Word doc\n"
        "`/clear` — Reset conversation\n"
        "`/stats` — Your usage stats\n"
        "`/model` — Active AI models\n"
        "`/help` — Tips & examples\n\n"
        "_Powered by Groq AI — Built by Richard_"
    )

def handle_help(chat_id):
    send_message(chat_id,
        "💡 *JARVIS Tips & Tricks*\n\n"
        "📄 *Make documents:*\n"
        "• `/pdf Python decorators`\n"
        "• `/doc Poultry farm business plan`\n"
        "• `make a pdf about Supabase`\n"
        "• `create word doc on M-Pesa Daraja API`\n\n"
        "🖼 *Analyze images:*\n"
        "• Send any screenshot or photo\n"
        "• Add a caption for specific questions\n"
        "• Works for code screenshots, errors, diagrams\n\n"
        "💻 *Code help:*\n"
        "• Paste code and ask questions\n"
        "• `Debug this: [your code]`\n"
        "• `Explain this PHP function`\n"
        "• `Write a Python script that...`\n\n"
        "🧠 Context memory: last 20 messages\n"
        "Use /clear to start fresh."
    )

def handle_stats(chat_id, user_name):
    stats       = user_stats.get(chat_id, {})
    msgs        = stats.get("messages", 0)
    since       = stats.get("first_seen", "today")
    history_len = len(get_history(chat_id))
    send_message(chat_id,
        f"📊 *Your JARVIS Stats*\n\n"
        f"👤 Name: {user_name}\n"
        f"💬 Messages sent: {msgs}\n"
        f"📅 Using since: {since}\n"
        f"🧠 Memory: {history_len}/20 messages\n"
        f"🤖 Text model: `{GROQ_MODEL}`\n"
        f"🔭 Vision model: `{GROQ_MODEL_VISION}`"
    )

def handle_clear(chat_id):
    conversation_history[chat_id] = []
    send_message(chat_id, "🗑️ Conversation cleared. Fresh start!")

def handle_model(chat_id):
    send_message(chat_id,
        f"🤖 *Active Models*\n\n"
        f"💬 Text: `{GROQ_MODEL}`\n"
        f"🔭 Vision: `{GROQ_MODEL_VISION}`\n"
        f"📦 Max tokens: `{MAX_TOKENS}`\n"
        f"📚 Memory depth: 20 messages"
    )

# ══════════════════════════════════════════════════════════════
#  MESSAGE PROCESSOR
# ══════════════════════════════════════════════════════════════
def process_message(message):
    chat_id   = message["chat"]["id"]
    user_name = message["from"].get("first_name", "User")
    text      = message.get("text", "").strip()
    photo     = message.get("photo")
    caption   = message.get("caption", "").strip()
    document  = message.get("document")

    track_usage(chat_id, user_name)
    log.info(f"[{user_name}:{chat_id}] {'[IMG]' if photo else '[DOC]' if document else text[:60]}")

    # ── Image ──────────────────────────────────────────────────
    if photo:
        send_typing(chat_id)
        send_message(chat_id, "🔍 Analyzing your image...")
        best = max(photo, key=lambda p: p.get("file_size", 0))
        b64  = download_telegram_file(best["file_id"])
        if b64:
            prompt = caption if caption else (
                "Analyze this image in detail. Describe everything you see "
                "and provide relevant insights or help."
            )
            reply = call_groq_vision(chat_id, b64, prompt)
            send_message(chat_id, reply)
        else:
            send_message(chat_id, "❌ Could not download image. Try again.")
        return

    # ── Document file ──────────────────────────────────────────
    if document:
        name = document.get("file_name", "file")
        send_message(chat_id,
            f"📎 Received file: *{name}*\n\n"
            "Document text extraction coming soon!\n"
            "For now, paste the text content directly.")
        return

    # ── No text ────────────────────────────────────────────────
    if not text:
        send_message(chat_id, "Send me a message or image! Type /help for what I can do.")
        return

    # ── Slash commands ─────────────────────────────────────────
    if text == "/start":
        handle_start(chat_id, user_name)
    elif text == "/help":
        handle_help(chat_id)
    elif text == "/clear":
        handle_clear(chat_id)
    elif text == "/model":
        handle_model(chat_id)
    elif text.startswith("/stats"):
        handle_stats(chat_id, user_name)

    elif text.lower().startswith("/pdf"):
        topic = text[4:].strip()
        if not topic:
            send_message(chat_id, "📄 Specify a topic.\nExample: `/pdf Python decorators`")
        else:
            handle_generate_doc(chat_id, topic, "pdf")

    elif text.lower().startswith("/doc"):
        topic = text[4:].strip()
        if not topic:
            send_message(chat_id, "📝 Specify a topic.\nExample: `/doc M-Pesa integration`")
        else:
            handle_generate_doc(chat_id, topic, "docx")

    # ── Natural language PDF triggers ──────────────────────────
    elif re.search(r'(make|create|generate|write|build)\s+(a\s+)?pdf', text, re.I):
        topic = re.sub(
            r'.*(make|create|generate|write|build)\s+(a\s+)?pdf\s*(about|on|for|of)?\s*',
            '', text, flags=re.IGNORECASE
        ).strip() or "General Topic"
        handle_generate_doc(chat_id, topic, "pdf")

    # ── Natural language DOC triggers ──────────────────────────
    elif re.search(r'(make|create|generate|write|build)\s+(a\s+)?(doc|word)', text, re.I):
        topic = re.sub(
            r'.*(make|create|generate|write|build)\s+(a\s+)?(doc|word\s*doc|document|word)\s*(about|on|for|of)?\s*',
            '', text, flags=re.IGNORECASE
        ).strip() or "General Topic"
        handle_generate_doc(chat_id, topic, "docx")

    # ── Regular chat ───────────────────────────────────────────
    else:
        send_typing(chat_id)
        add_to_history(chat_id, "user", text)
        reply = ask_ai(text)
        add_to_history(chat_id, "assistant", reply)
        send_message(chat_id, reply)

# ══════════════════════════════════════════════════════════════
#  MAIN LOOP
# ══════════════════════════════════════════════════════════════
def main():
    log.info("╔═══════════════════════════════════╗")
    log.info("║    KILO JARVIS v2.0 — Starting    ║")
    log.info("╚═══════════════════════════════════╝")
    log.info(f"Text Model  : {GROQ_MODEL}")
    log.info(f"Vision Model: {GROQ_MODEL_VISION}")
    log.info(f"PDF  Support: {'✅ Ready' if PDF_AVAILABLE  else '❌  pip install reportlab'}")
    log.info(f"DOCX Support: {'✅ Ready' if DOCX_AVAILABLE else '❌  pip install python-docx'}")
    log.info("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")

    offset = None
    while True:
        updates = get_updates(offset)
        for update in updates:
            offset = update["update_id"] + 1
            if "message" in update:
                try:
                    process_message(update["message"])
                except Exception as e:
                    log.error(f"Process error: {e}", exc_info=True)
        if not updates:
            time.sleep(0.3)

if __name__ == "__main__":
    main()
